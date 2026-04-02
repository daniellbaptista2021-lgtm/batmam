"""Gerador de documentos DOCX."""
from __future__ import annotations
import json
from pathlib import Path
from .base import STATIC_DIR, ask_ai, unique_name, file_url

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate(prompt: str) -> dict:
    system = """Voce e um redator profissional. Gere um documento estruturado em JSON.
Formato EXATO (sem markdown, sem explicacao):
{
  "title": "Titulo do Documento",
  "sections": [
    {
      "heading": "Titulo da Secao",
      "level": 1,
      "paragraphs": ["Texto do paragrafo 1.", "Texto do paragrafo 2."]
    },
    {
      "heading": "Subsecao",
      "level": 2,
      "paragraphs": ["Texto..."]
    }
  ]
}
- Conteudo profissional e completo em portugues brasileiro
- Minimo 3 secoes com paragrafos detalhados
- Retorne APENAS o JSON valido"""

    raw = ask_ai(prompt, system=system, max_tokens=4096)

    text = raw.strip()
    if text.startswith("```"):
        text = "\n".join(text.split("\n")[1:])
    if text.endswith("```"):
        text = text[:-3]
    text = text.strip()

    data = json.loads(text)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    title_para = doc.add_heading(data.get("title", "Documento"), level=0)
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

    for section in data.get("sections", []):
        level = min(section.get("level", 1), 3)
        heading = doc.add_heading(section.get("heading", ""), level=level)
        for run in heading.runs:
            if level == 1:
                run.font.color.rgb = RGBColor(0x5B, 0x21, 0xB6)

        for para_text in section.get("paragraphs", []):
            p = doc.add_paragraph(para_text)
            p.paragraph_format.space_after = Pt(6)

    filename = unique_name(data.get("title", prompt[:30]), ".docx")
    out_dir = STATIC_DIR / "files"
    out_dir.mkdir(parents=True, exist_ok=True)
    filepath = out_dir / filename
    doc.save(filepath)

    url = file_url(f"static/files/{filename}")
    size = filepath.stat().st_size

    return {
        "type": "docx",
        "name": filename,
        "url": url,
        "size": size,
        "path": str(filepath),
    }
